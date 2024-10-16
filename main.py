import os
from flask import Flask, render_template, send_file, request, abort
from flask_wtf import FlaskForm
from wtforms import StringField, DateField, SubmitField, DecimalField
from wtforms.validators import DataRequired, Optional
from docx import Document
import io
from datetime import date, datetime
import logging
import mimetypes

# Get the directory of the current file (app.py)
basedir = os.path.abspath(os.path.dirname(__file__))

app = Flask(__name__)
app.config['SECRET_KEY'] = 'your_secret_key_here'
app.config['TEMPLATE_PATH'] = os.path.join(basedir, 'documents', 'plantilla.docx')
logging.basicConfig(level=logging.DEBUG)

mimetypes.add_type('application/vnd.openxmlformats-officedocument.wordprocessingml.document', '.docx')

def check_write_permissions(path):
    if os.access(path, os.W_OK):
        app.logger.info(f"Write permissions OK for {path}")
    else:
        app.logger.error(f"No write permissions for {path}")

class FormularioEnvio(FlaskForm):
    # Exportador fields
    ruc_exportador = StringField('RUC Exportador', validators=[DataRequired()])
    razon_social_exportador = StringField('Razón Social Exportador', validators=[DataRequired()])
    domicilio_fiscal_exportador = StringField('Domicilio Fiscal Exportador', validators=[DataRequired()])
    locales_anexos_exportador = StringField('Locales Anexos Exportador', validators=[Optional()])
    objeto_social_exportador = StringField('Objeto Social Exportador', validators=[DataRequired()])
    representante_legal_exportador = StringField('Representante Legal Exportador', validators=[DataRequired()])
    dni_exportador = StringField('DNI Exportador', validators=[DataRequired()])
    cargo_exportador = StringField('Cargo Exportador', validators=[DataRequired()])
    telefono_exportador = StringField('Teléfono Exportador', validators=[DataRequired()])
    origen_de_los_fondos_exportador = StringField('Origen de los Fondos Exportador', validators=[DataRequired()])
    consignee_importador_exportador = StringField('Consignee/Importador Exportador', validators=[DataRequired()])
    domicilio_fiscal_consignee_exportador = StringField('Domicilio Fiscal Consignee Exportador', validators=[DataRequired()])
    telefono_consignee_exportador = StringField('Teléfono Consignee Exportador', validators=[DataRequired()])
    email_consignee_exportador = StringField('Email Consignee Exportador', validators=[DataRequired()])
    attn_consignee_exportador = StringField('Attn Consignee Exportador', validators=[DataRequired()])
    notify1_exportador = StringField('Notify1 Exportador', validators=[DataRequired()])
    ruc_notify_exportador = StringField('RUC Notify Exportador', validators=[DataRequired()])
    domicilio_fiscal_notify_exportador = StringField('Domicilio Fiscal Notify Exportador', validators=[DataRequired()])
    notify2_exportador = StringField('Notify2 Exportador', validators=[Optional()])
    ruc_notify2_exportador = StringField('RUC Notify2 Exportador', validators=[Optional()])
    domicilio_fiscal_notify2_exportador = StringField('Domicilio Fiscal Notify2 Exportador', validators=[Optional()])
    cliente_empresa_exportador = StringField('Cliente Empresa Exportador', validators=[DataRequired()])
    destino_exportador = StringField('Destino Exportador', validators=[DataRequired()])
    peso_kg_exportador = DecimalField('Peso KG Exportador', validators=[DataRequired()])
    fecha_de_ingreso_exportador = DateField('Fecha de Ingreso Exportador', format='%Y-%m-%d', validators=[DataRequired()])
    fecha_de_embarque_exportador = DateField('Fecha de Embarque Exportador', format='%Y-%m-%d', validators=[DataRequired()])
    guia_aerea_exportador = StringField('Guía Aérea Exportador', validators=[DataRequired()])
    aerolinea_exportador = StringField('Aerolínea Exportador', validators=[DataRequired()])
    invoice_origen_exportador = StringField('Invoice Origen Exportador', validators=[DataRequired()])
    mining_code_exportador = StringField('Mining Code Exportador', validators=[DataRequired()])
    guia_de_remision_exportador = StringField('Guía de Remisión Exportador', validators=[DataRequired()])
    reinpo_exportador = StringField('REINPO Exportador', validators=[DataRequired()])
    ciudad_exportador = StringField('Ciudad Exportador', validators=[DataRequired()])
    fecharegistro_exportador = DateField('Fecha Registro Exportador', format='%Y-%m-%d', validators=[DataRequired()])

    # Productor fields
    ruc_productor = StringField('RUC Productor', validators=[DataRequired()])
    razon_social_productor = StringField('Razón Social Productor', validators=[DataRequired()])
    domicilio_fiscal_productor = StringField('Domicilio Fiscal Productor', validators=[DataRequired()])
    dni_productor = StringField('DNI Productor', validators=[DataRequired()])
    departamento_productor = StringField('Departamento Productor', validators=[DataRequired()])
    factura_productor = StringField('Factura N° Productor', validators=[DataRequired()])
    fecha_factura_productor = DateField('Fecha Factura Productor', format='%Y-%m-%d', validators=[DataRequired()])
    valor_fob_productor = DecimalField('Valor FOB Productor', validators=[DataRequired()])
    peso_neto_productor = DecimalField('Peso Neto Productor', validators=[DataRequired()])
    peso_fino_productor = DecimalField('Peso Fino Productor', validators=[DataRequired()])
    tipo_producto_productor = StringField('Tipo de Producto Productor', validators=[DataRequired()])
    barra_de_oro_dore_productor = StringField('Barra de Oro Doré Productor', validators=[DataRequired()])
    peso_no_oro_productor = DecimalField('Peso No Oro Productor', validators=[DataRequired()])
    guia_remision_productor = StringField('Guía de Remisión Productor', validators=[DataRequired()])
    recpo_n_productor = StringField('RECPO N° Productor', validators=[DataRequired()])
    concesion_minera_productor = StringField('Concesión Minera Productor', validators=[DataRequired()])
    codigo_ingemmet_n_productor = StringField('Código INGEMMET N° Productor', validators=[DataRequired()])
    fecha_registro_productor = DateField('Fecha Registro Productor', format='%Y-%m-%d', validators=[DataRequired()])
    fecha_guia_remision_productor = DateField('Fecha Guía Remisión Productor', format='%Y-%m-%d', validators=[DataRequired()])

    # Transporte fields
    ruc_transporte = StringField('RUC Transporte', validators=[DataRequired()])
    razon_social_transporte = StringField('Razón Social Transporte', validators=[DataRequired()])
    domicilio_fiscal_transporte = StringField('Domicilio Fiscal Transporte', validators=[DataRequired()])
    vehiculo_transporte = StringField('Vehículo Transporte', validators=[DataRequired()])
    marca_transporte = StringField('Marca Transporte', validators=[DataRequired()])
    placa_transporte = StringField('Placa Transporte', validators=[DataRequired()])
    guia_remision_transporte = StringField('Guía de Remisión Transportista', validators=[DataRequired()])

    submit = SubmitField('Enviar')

@app.route('/', methods=['GET', 'POST'])
def index():
    form = FormularioEnvio()
    if request.method == 'POST':
        app.logger.info("Form submitted")
        app.logger.info(f"Form data: {request.form}")
        if form.validate():
            app.logger.info("Form validated successfully")
        else:
            app.logger.error(f"Form validation failed: {form.errors}")
            for field, errors in form.errors.items():
                app.logger.error(f"Field {field} errors: {errors}")
    
    if form.validate_on_submit():
        try:
            app.logger.info("Form validated on submit")
            doc = generar_documento(form)
            
            output_dir = os.path.join(basedir, 'output')
            app.logger.info(f"Output directory: {output_dir}")
            if not os.path.exists(output_dir):
                os.makedirs(output_dir)
                app.logger.info(f"Created output directory: {output_dir}")
            
            check_write_permissions(output_dir)
            
            output_path = os.path.join(output_dir, 'instruccion_de_envio.docx')
            app.logger.info(f"Saving document to: {output_path}")
            doc.save(output_path)
            app.logger.info(f"Document saved successfully to {output_path}")
            
            if os.path.exists(output_path):
                file_size = os.path.getsize(output_path)
                app.logger.info(f"File exists. Size: {file_size} bytes")
            else:
                app.logger.error(f"File does not exist at {output_path}")
            
            try:
                with open(output_path, 'rb') as f:
                    content = f.read()
                app.logger.info(f"Successfully read {len(content)} bytes from the file")
            except Exception as e:
                app.logger.error(f"Error reading saved file: {str(e)}")
            
            app.logger.info("Attempting to send file")
            return send_file(
                output_path,
                as_attachment=True,
                download_name='instruccion_de_envio.docx',
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
        except Exception as e:
            app.logger.error(f"Error in index route: {str(e)}")
            return f"An error occurred: {str(e)}", 500
    
    return render_template('index.html', form=form)

@app.errorhandler(500)
def internal_server_error(e):
    return "Internal Server Error: " + str(e), 500

@app.errorhandler(Exception)
def handle_exception(e):
    app.logger.error(f"Unhandled exception: {str(e)}")
    return "An unexpected error occurred", 500

def generar_documento(form):
    try:
        app.logger.info(f"Template path: {app.config['TEMPLATE_PATH']}")
        doc = Document(app.config['TEMPLATE_PATH'])
        app.logger.info("Document loaded from template")
        
        replacements = {
            # Exportador replacements
            '[RUC_EXPORTADOR]': form.ruc_exportador.data,
            '[RAZON_SOCIAL_EXPORTADOR]': form.razon_social_exportador.data,
            '[DOMICILIO_FISCAL_EXPORTADOR]': form.domicilio_fiscal_exportador.data,
            '[LOCALES_ANEXOS_EXPORTADOR]': form.locales_anexos_exportador.data,
            '[OBJETO_SOCIAL_EXPORTADOR]': form.objeto_social_exportador.data,
            '[REPRESENTANTE_LEGAL_EXPORTADOR]': form.representante_legal_exportador.data,
            '[DNI_EXPORTADOR]': form.dni_exportador.data,
            '[CARGO_EXPORTADOR]': form.cargo_exportador.data,
            '[TELEFONO_EXPORTADOR]': form.telefono_exportador.data,
            '[ORIGEN_DE_LOS_FONDOS_EXPORTADOR]': form.origen_de_los_fondos_exportador.data,
            '[CONSIGNEE_IMPORTADOR_EXPORTADOR]': form.consignee_importador_exportador.data,
            '[DOMICILIO_FISCAL_CONSIGNEE_EXPORTADOR]': form.domicilio_fiscal_consignee_exportador.data,
            '[TELEFONO_CONSIGNEE_EXPORTADOR]': form.telefono_consignee_exportador.data,
            '[EMAIL_CONSIGNEE_EXPORTADOR]': form.email_consignee_exportador.data,
            '[ATTN_CONSIGNEE_EXPORTADOR]': form.attn_consignee_exportador.data,
            '[NOTIFY1_EXPORTADOR]': form.notify1_exportador.data,
            '[RUC_NOTIFY_EXPORTADOR]': form.ruc_notify_exportador.data,
            '[DOMICILIO_FISCAL_NOTIFY_EXPORTADOR]': form.domicilio_fiscal_notify_exportador.data,
            '[NOTIFY2_EXPORTADOR]': form.notify2_exportador.data,
            '[RUC_NOTIFY2_EXPORTADOR]': form.ruc_notify2_exportador.data,
            '[DOMICILIO_FISCAL_NOTIFY2_EXPORTADOR]': form.domicilio_fiscal_notify2_exportador.data,
            '[CLIENTE_EMPRESA_EXPORTADOR]': form.cliente_empresa_exportador.data,
            '[DESTINO_EXPORTADOR]': form.destino_exportador.data,
            '[PESO_KG_EXPORTADOR]': str(form.peso_kg_exportador.data),
            '[FECHA_DE_INGRESO_EXPORTADOR]': form.fecha_de_ingreso_exportador.data.strftime('%d/%m/%Y') if form.fecha_de_ingreso_exportador.data else '',
            '[FECHA_DE_EMBARQUE_EXPORTADOR]': form.fecha_de_embarque_exportador.data.strftime('%d/%m/%Y') if form.fecha_de_embarque_exportador.data else '',
            '[GUIA_AEREA_EXPORTADOR]': form.guia_aerea_exportador.data,
            '[AEROLINEA_EXPORTADOR]': form.aerolinea_exportador.data,
            '[INVOICE_ORIGEN_EXPORTADOR]': form.invoice_origen_exportador.data,
            '[MINING_CODE_EXPORTADOR]': form.mining_code_exportador.data,
            '[GUIA_DE_REMISION_EXPORTADOR]': form.guia_de_remision_exportador.data,
            '[REINPO_EXPORTADOR]': form.reinpo_exportador.data,
            '[CIUDAD_EXPORTADOR]': form.ciudad_exportador.data,
            '[FECHAREGISTRO_EXPORTADOR]': form.fecharegistro_exportador.data.strftime('%d/%m/%Y') if form.fecharegistro_exportador.data else '',
            
            # Productor replacements
            '[RUC_PRODUCTOR]': form.ruc_productor.data,
            '[RAZON_SOCIAL_PRODUCTOR]': form.razon_social_productor.data,
            '[DOMICILIO_FISCAL_PRODUCTOR]': form.domicilio_fiscal_productor.data,
            '[DNI_PRODUCTOR]': form.dni_productor.data,
            '[DEPARTAMENTO_PRODUCTOR]': form.departamento_productor.data,
            '[FACTURA_PRODUCTOR]': form.factura_productor.data,
            '[FECHA_FACTURA_PRODUCTOR]': form.fecha_factura_productor.data.strftime('%d/%m/%Y') if form.fecha_factura_productor.data else '',
            '[VALOR_FOB_PRODUCTOR]': str(form.valor_fob_productor.data),
            '[PESO_NETO_PRODUCTOR]': str(form.peso_neto_productor.data),
            '[PESO_FINO_PRODUCTOR]': str(form.peso_fino_productor.data),
            '[TIPO_PRODUCTO_PRODUCTOR]': form.tipo_producto_productor.data,
            '[BARRA_DE_ORO_DORE_PRODUCTOR]': form.barra_de_oro_dore_productor.data,
            '[PESO_NO_ORO_PRODUCTOR]': str(form.peso_no_oro_productor.data),
            '[GUIA_REMISION_PRODUCTOR]': form.guia_remision_productor.data,
            '[RECPO_N_PRODUCTOR]': form.recpo_n_productor.data,
            '[CONCESION_MINERA_PRODUCTOR]': form.concesion_minera_productor.data,
            '[CODIGO_INGEMMET_N_PRODUCTOR]': form.codigo_ingemmet_n_productor.data,
            '[FECHA_REGISTRO_PRODUCTOR]': form.fecha_registro_productor.data.strftime('%d/%m/%Y') if form.fecha_registro_productor.data else '',
            '[FECHA_GUIA_REMISION_PRODUCTOR]': form.fecha_guia_remision_productor.data.strftime('%d/%m/%Y') if form.fecha_guia_remision_productor.data else '',
            
            # Transporte replacements
            '[RUC_TRANSPORTE]': form.ruc_transporte.data,
            '[RAZON_SOCIAL_TRANSPORTE]': form.razon_social_transporte.data,
            '[DOMICILIO_FISCAL_TRANSPORTE]': form.domicilio_fiscal_transporte.data,
            '[VEHICULO_TRANSPORTE]': form.vehiculo_transporte.data,
            '[MARCA_TRANSPORTE]': form.marca_transporte.data,
            '[PLACA_TRANSPORTE]': form.placa_transporte.data,
            '[GUIA_REMISION_TRANSPORTE]': form.guia_remision_transporte.data,
        }
        app.logger.info("Replacements dictionary created")

        # Function to replace text in paragraphs and tables
        def replace_text(obj):
            if hasattr(obj, 'paragraphs'):
                for paragraph in obj.paragraphs:
                    for key, value in replacements.items():
                        if key in paragraph.text:
                            paragraph.text = paragraph.text.replace(key, str(value))
            if hasattr(obj, 'tables'):
                for table in obj.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for key, value in replacements.items():
                                if key in cell.text:
                                    cell.text = cell.text.replace(key, str(value))
            app.logger.info(f"Replaced text in {type(obj).__name__}")

        # Replace text in the main document body
        replace_text(doc)
        app.logger.info("Replaced text in main document body")

        # Replace text in headers and footers
        for section in doc.sections:
            replace_text(section.header)
            replace_text(section.footer)
        app.logger.info("Replaced text in headers and footers")

        app.logger.info("Document generation completed")
        return doc
    except Exception as e:
        app.logger.error(f"Error in generar_documento: {str(e)}")
        raise

if __name__ == '__main__':
    app.run(debug=True)
